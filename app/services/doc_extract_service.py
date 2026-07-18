"""Extract text from uploaded documents/images and classify the kind."""

from __future__ import annotations

import logging
from pathlib import Path

from app.config import get_settings

logger = logging.getLogger(__name__)

IMAGE_MIMES = {"image/jpeg", "image/png", "image/webp", "image/heic", "image/heif"}


def extract_text(file_path: str, mime_type: str, file_name: str = "") -> tuple[str, str]:
    """Return (source_type, extracted_text)."""
    name = (file_name or file_path).lower()
    mime = (mime_type or "").lower()

    if mime == "application/pdf" or name.endswith(".pdf"):
        return "pdf", _extract_pdf(file_path)
    if name.endswith(".docx") or "word" in mime:
        return "docx", _extract_docx(file_path)
    if mime in IMAGE_MIMES or name.endswith((".jpg", ".jpeg", ".png", ".webp")):
        return "image", _extract_image(file_path, mime or "image/jpeg")
    # Fallback: try reading as text
    try:
        return "other", Path(file_path).read_text(encoding="utf-8", errors="ignore")[:20000]
    except Exception:
        return "other", ""


def _extract_pdf(file_path: str) -> str:
    text = ""
    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            parts = [(page.extract_text() or "") for page in pdf.pages]
        text = "\n".join(parts).strip()
    except Exception as e:
        logger.warning("pdfplumber_failed err=%s", e)

    if len(text) < 20:
        # Likely scanned — use Gemini vision on rendered text via pypdf fallback
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            text = "\n".join((p.extract_text() or "") for p in reader.pages).strip()
        except Exception as e:
            logger.warning("pypdf_failed err=%s", e)

    if len(text) < 20:
        logger.info("pdf_text_empty — trying gemini vision")
        text = _gemini_file(file_path, "application/pdf")
    logger.info("pdf_extract chars=%s", len(text))
    return text


def _extract_docx(file_path: str) -> str:
    try:
        import docx

        d = docx.Document(file_path)
        text = "\n".join(p.text for p in d.paragraphs)
        for table in d.tables:
            for row in table.rows:
                text += "\n" + " | ".join(c.text for c in row.cells)
        logger.info("docx_extract chars=%s", len(text))
        return text.strip()
    except Exception as e:
        logger.warning("docx_failed err=%s", e)
        return ""


def _extract_image(file_path: str, mime_type: str) -> str:
    return _gemini_file(file_path, mime_type)


def _gemini_file(file_path: str, mime_type: str) -> str:
    from google import genai
    from google.genai import types

    settings = get_settings()
    if not settings.gemini_api_key:
        return ""
    client = genai.Client(api_key=settings.gemini_api_key)
    data = Path(file_path).read_bytes()
    prompt = (
        "Read this document. It may be a supplier invoice, purchase list, or handwritten "
        "grocery order. Extract ALL line items with product name, quantity, unit, cost price, "
        "and MRP if present. Return as clean readable text (one item per line). "
        "Also state at the top what kind of document this looks like."
    )
    resp = client.models.generate_content(
        model=settings.gemini_model,
        contents=[
            types.Content(
                role="user",
                parts=[
                    types.Part.from_text(text=prompt),
                    types.Part.from_bytes(data=data, mime_type=mime_type),
                ],
            )
        ],
    )
    return (resp.text or "").strip()


def quick_classify(text: str) -> str:
    """Cheap heuristic; the agent still reasons over it."""
    t = (text or "").lower()
    if "tax invoice" in t or ("gstin" in t and "hsn" in t):
        return "supplier_invoice"
    if "gstr" in t or "gst return" in t:
        return "gst_return"
    if "statement" in t and ("upi" in t or "bank" in t or "a/c" in t):
        return "bank_statement"
    if any(k in t for k in ("qty", "quantity", "pcs", "packet", "kg", "x ")):
        return "purchase_list"
    return "unknown"
